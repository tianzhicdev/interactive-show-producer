#!/usr/bin/env python3
"""Render interactive show project state into PDF + DOCX deliverables.

Usage:
    python renderer.py <project_dir>

Reads:  <project_dir>/state.json
Writes: <project_dir>/output/互动剧本_<title>.pdf
        <project_dir>/output/互动剧本_<title>_EP01.docx
        ...
"""
import json
import os
import re
import sys
import tempfile
import subprocess

# ── PDF rendering via fpdf2 ──────────────────────────────────────────

def find_cjk_font():
    """Find a CJK-capable TTF/TTC font on the system."""
    candidates = [
        '/System/Library/Fonts/Hiragino Sans GB.ttc',
        '/System/Library/Fonts/STHeiti Medium.ttc',
        '/System/Library/Fonts/PingFang.ttc',
        '/Library/Fonts/Arial Unicode.ttf',
        '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
        # Linux
        '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def normalize_fork_points(structure):
    """Convert legacy fork_structure to fork_points array, or return existing fork_points."""
    if "fork_points" in structure:
        return structure["fork_points"]

    fork = structure.get("fork_structure", {})
    if not fork:
        return []

    threads = {}
    for name, info in fork.get("parallel_threads", {}).items():
        threads[name] = {"episodes": info.get("episodes", []), "theme": info.get("theme", "")}

    return [{
        "id": "FORK_1",
        "fork_choice": fork.get("fork_point", ""),
        "threads": threads,
        "convergence_episode": fork.get("convergence_episode")
    }]


def render_pdf(project_dir: str, state: dict, preview: bool = False):
    """Render the full deliverable PDF: summary → graph → scripts.

    If preview=True, render only cover + summary + stats + overview graph (no scripts, no per-episode graphs).
    """
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    meta = state['metadata']
    title = meta.get('title', '互动影游')
    output_dir = os.path.join(project_dir, 'output')
    os.makedirs(output_dir, exist_ok=True)

    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=20)

    # Register CJK font (no bold variant — use regular for both)
    font_path = find_cjk_font()
    if font_path:
        pdf.add_font('CJK', '', font_path)
        body_font = 'CJK'
    else:
        body_font = 'Helvetica'
        print('WARNING: No CJK font found. Chinese characters may not render.')

    # Helpers: always reset x to left margin after output
    def cell_ln(w, h, txt, **kwargs):
        pdf.cell(w, h, txt, new_x=XPos.LMARGIN, new_y=YPos.NEXT, **kwargs)

    def mcell(w, h, txt, **kwargs):
        pdf.multi_cell(w, h, txt, new_x=XPos.LMARGIN, new_y=YPos.NEXT, **kwargs)

    # ── Cover page ──
    pdf.add_page()
    pdf.set_font(body_font, size=28)
    pdf.ln(60)
    cell_ln(0, 15, title, align='C')
    pdf.set_font(body_font, size=14)
    pdf.ln(10)
    cell_ln(0, 10, '互动影游剧本', align='C')
    pdf.ln(5)
    # Compute durations from episode data
    _struct = state.get('structure', {})
    episodes_list = _struct.get('episodes', [])
    total_ep_minutes = sum(ep.get('duration_minutes', 0) for ep in episodes_list)

    # Per-episode duration: from episodes or from metadata (e.g., "3m" → 3)
    per_ep_dur = episodes_list[0].get('duration_minutes', 0) if episodes_list else 0
    if not per_ep_dur:
        dur_str = meta.get('episode_duration', '3m')
        try:
            per_ep_dur = int(re.sub(r'[^0-9]', '', str(dur_str)) or '3')
        except ValueError:
            per_ep_dur = 3
        total_ep_minutes = len(episodes_list) * per_ep_dur

    # Playthrough episodes: try stats, meta_stats, episode_graph, then derive from ratio
    _stats = _struct.get('stats', {})
    eps_per_play = _stats.get('playthrough_episodes', 0)
    if not eps_per_play:
        eps_per_play = _struct.get('meta_stats', {}).get('episodes_per_playthrough', 0)
    if not eps_per_play:
        eps_per_play = _struct.get('episode_graph', {}).get('single_playthrough_episodes', 0)
    if not eps_per_play:
        ratio = meta.get('playthrough_ratio', 0.6)
        if isinstance(ratio, str):
            try:
                ratio = float(ratio)
            except ValueError:
                ratio = 0.6
        eps_per_play = round(len(episodes_list) * ratio) if episodes_list else 0
    playthrough_minutes = eps_per_play * per_ep_dur if eps_per_play and per_ep_dur else None

    playthrough_str = f"{playthrough_minutes}分钟" if playthrough_minutes else meta.get('playthrough_duration', 'N/A')
    total_str = f"{total_ep_minutes}分钟" if total_ep_minutes else meta.get('total_duration', 'N/A')

    # Compute choices per episode from structure data (supports both 'choices' array and 'choice' dict)
    total_choices = 0
    for ep in episodes_list:
        if ep.get('choices'):
            total_choices += len(ep['choices'])
        elif ep.get('choice') and isinstance(ep['choice'], dict):
            total_choices += 1
    avg_choices = round(total_choices / len(episodes_list), 1) if episodes_list else 0
    selections_str = meta.get('selections_per_episode') or (str(avg_choices) if avg_choices else 'N/A')

    info_lines = [
        f"单次体验时长: {playthrough_str}（{eps_per_play}集 × {per_ep_dur}分钟）",
        f"全部集总时长: {total_str}（共{len(episodes_list)}集）",
        f"每集互动节点: {selections_str}",
        f"每节点选项数: {meta.get('options_per_selection', 'N/A')}",
    ]
    pdf.set_font(body_font, size=11)
    for line in info_lines:
        cell_ln(0, 8, line, align='C')
    if meta.get('note'):
        pdf.ln(10)
        pdf.set_font(body_font, size=9)
        mcell(0, 6, f"创作方向: {meta['note']}")

    # ── Section 1: Summary / Story Bible ──
    pdf.add_page()
    pdf.set_font(body_font, size=18)
    cell_ln(0, 12, '◆ 第一部分: 故事摘要')
    pdf.ln(5)

    bible = state.get('story_bible', {})
    if bible.get('one_line_summary'):
        pdf.set_font(body_font, size=12)
        cell_ln(0, 8, '● 一句话概括')
        pdf.set_font(body_font, size=11)
        mcell(0, 7, bible['one_line_summary'])
        pdf.ln(5)

    if bible.get('world_settings'):
        pdf.set_font(body_font, size=12)
        cell_ln(0, 8, '● 世界观设定')
        pdf.set_font(body_font, size=10)
        ws = bible['world_settings']
        if isinstance(ws, dict):
            import json as _json
            ws_text = '\n'.join(f'{k}: {v if isinstance(v, str) else _json.dumps(v, ensure_ascii=False)}' for k, v in ws.items())
        else:
            ws_text = str(ws)
        mcell(0, 6, ws_text)
        pdf.ln(5)

    if bible.get('characters'):
        pdf.set_font(body_font, size=12)
        cell_ln(0, 8, '● 核心角色')
        pdf.set_font(body_font, size=10)
        for char in bible['characters']:
            if isinstance(char, dict):
                name = char.get('name', '')
                role = char.get('role', '')
                desc = char.get('description', '')
                pdf.set_font(body_font, size=10)
                cell_ln(0, 7, f"【{name}】{role}")
                pdf.set_font(body_font, size=10)
                mcell(0, 6, desc)
                pdf.ln(2)
            else:
                mcell(0, 6, str(char))
        pdf.ln(5)

    if bible.get('plot_summary'):
        pdf.set_font(body_font, size=12)
        cell_ln(0, 8, '● 剧情概要')
        pdf.set_font(body_font, size=10)
        mcell(0, 6, bible['plot_summary'])
        pdf.ln(5)

    # Episode summaries
    structure = state.get('structure', {})
    episodes = structure.get('episodes', [])
    if episodes:
        pdf.set_font(body_font, size=12)
        cell_ln(0, 8, '● 分集概要')
        pdf.ln(3)
        for ep in episodes:
            pdf.set_font(body_font, size=11)
            ep_id = ep.get('id', '')
            ep_title = ep.get('title', '')
            cell_ln(0, 7, f"【{ep_id}】{ep_title}")
            pdf.set_font(body_font, size=10)
            if ep.get('summary'):
                mcell(0, 6, ep['summary'])
            pdf.ln(3)

    # ── Section 1.5: Player Stats Overview (preview + full) ──
    player_stats = structure.get('player_stats', [])
    if player_stats:
        pdf.set_font(body_font, size=12)
        cell_ln(0, 8, '● 数值系统')
        pdf.ln(3)
        for ps in player_stats:
            pdf.set_font(body_font, size=10)
            name = ps.get('name', ps.get('id', ''))
            desc = ps.get('description', '')
            init = ps.get('initial_value', 50)
            thresholds = ps.get('thresholds', {})
            low_t = thresholds.get('low', '?')
            high_t = thresholds.get('high', '?')
            cell_ln(0, 7, f"【{name}】初始值: {init}  低阈值: < {low_t}  高阈值: > {high_t}")
            if desc:
                pdf.set_font(body_font, size=9)
                mcell(0, 5, f"  {desc}")
            low_eff = ps.get('low_effect', '')
            high_eff = ps.get('high_effect', '')
            if low_eff:
                pdf.set_font(body_font, size=9)
                pdf.set_text_color(180, 0, 0)
                mcell(0, 5, f"  低值效果: {low_eff}")
                pdf.set_text_color(0, 0, 0)
            if high_eff:
                pdf.set_font(body_font, size=9)
                pdf.set_text_color(0, 120, 0)
                mcell(0, 5, f"  高值效果: {high_eff}")
                pdf.set_text_color(0, 0, 0)
            pdf.ln(2)
        pdf.ln(5)

    # ── Section 1.6: Fork Points / Branch Summary (preview mode) ──
    fork_points = normalize_fork_points(structure)
    # Lean format: derive fork summary from episodes where choice.type == "fork"
    if not fork_points and preview:
        lean_forks = []
        for ep in episodes:
            ch = ep.get('choice')
            if isinstance(ch, dict) and ch.get('type') == 'fork':
                targets = {}
                for opt in ch.get('options', []):
                    nxt = opt.get('next', '')
                    if nxt:
                        targets[opt.get('text', nxt)] = nxt
                if len(targets) >= 2:
                    lean_forks.append({'episode': ep['id'], 'question': ch.get('question', ''), 'targets': targets})
        if lean_forks:
            pdf.set_font(body_font, size=12)
            cell_ln(0, 8, '● 分支结构概览')
            pdf.ln(3)
            for lf in lean_forks:
                pdf.set_font(body_font, size=10)
                n_way = len(lf['targets'])
                cell_ln(0, 7, f"【{lf['episode']}】{n_way} 路分岔: {lf['question']}")
                pdf.set_font(body_font, size=9)
                for text, target in lf['targets'].items():
                    mcell(0, 5, f"  {text} → {target}")
                pdf.ln(2)
            # Stats
            _lean_stats = structure.get('stats', {})
            pp = _lean_stats.get('playthrough_episodes', 0)
            total_endings = _lean_stats.get('total_endings', len(structure.get('endings', [])))
            if pp:
                pdf.ln(2)
                pdf.set_font(body_font, size=10)
                cell_ln(0, 7, f"结局数: {total_endings}  单次通关集数: {pp}/{len(episodes)}")
            pdf.ln(5)
    if fork_points and preview:
        pdf.set_font(body_font, size=12)
        cell_ln(0, 8, '● 分支结构概览')
        pdf.ln(3)
        for fp in fork_points:
            pdf.set_font(body_font, size=10)
            threads = fp.get('threads', {})
            thread_names = list(threads.keys())
            n_way = len(thread_names)
            conv = fp.get('convergence_episode')
            conv_str = f" → 汇聚于 {conv}" if conv else " （终局分支）"
            cell_ln(0, 7, f"【{fp.get('id', '')}】{n_way} 路分岔{conv_str}")
            pdf.set_font(body_font, size=9)
            for tname, tinfo in threads.items():
                eps = ', '.join(tinfo.get('episodes', []))
                theme = tinfo.get('theme', '')
                theme_str = f" — {theme}" if theme else ''
                mcell(0, 5, f"  {tname}: [{eps}]{theme_str}")
            pdf.ln(2)
        # Overall stats
        ep_graph = structure.get('episode_graph', {})
        total_paths = ep_graph.get('total_unique_paths', 0)
        pp = structure.get('stats', {}).get('episodes_per_playthrough',
             ep_graph.get('single_playthrough_episodes', 0))
        if total_paths or pp:
            pdf.ln(2)
            pdf.set_font(body_font, size=10)
            cell_ln(0, 7, f"独立路线数: {total_paths}  单次通关集数: {pp}/{len(episodes)}")
        pdf.ln(5)

    # ── Section 2: Interactive Structure Graph ──
    # Render graphs via Graphviz (overview + per-episode)
    graph_images = render_graphviz_graphs(state, project_dir, preview=preview)
    first_graph = True
    for ep_id, img_path in graph_images:
        if not os.path.exists(img_path):
            continue

        # Read PNG dimensions from IHDR chunk (no PIL needed)
        import struct
        with open(img_path, 'rb') as _f:
            _f.read(16)
            img_px_w = struct.unpack('>I', _f.read(4))[0]
            img_px_h = struct.unpack('>I', _f.read(4))[0]
        ratio = img_px_w / img_px_h if img_px_h else 1.0

        # Choose orientation: landscape for wide images, portrait for tall
        if ratio > 1.0:
            pdf.add_page('L')
            page_w, page_h = 297, 210
            max_w, max_h = 277, 165  # margins: 10 left/right, 25 top, 20 bottom
            y_start = 25
        else:
            pdf.add_page('P')
            page_w, page_h = 210, 297
            max_w, max_h = 190, 255  # margins: 10 left/right, 25 top, 17 bottom
            y_start = 25

        # Section header on the first graph page only
        if first_graph:
            pdf.set_font(body_font, size=16)
            cell_ln(0, 10, '◆ 第二部分: 互动结构图', align='C')
            y_start += 12
            max_h -= 12
            first_graph = False

        # Title
        pdf.set_font(body_font, size=13)
        if ep_id == 'OVERVIEW':
            cell_ln(0, 9, '全剧互动结构总览', align='C')
        else:
            ep_info = next((e for e in episodes if e.get('id') == ep_id), {})
            cell_ln(0, 9, f"【{ep_id}】{ep_info.get('title', '')}", align='C')

        # Fit image within available area, preserving aspect ratio
        fit_w = max_w
        fit_h = fit_w / ratio
        if fit_h > max_h:
            fit_h = max_h
            fit_w = fit_h * ratio
        x_pos = (page_w - fit_w) / 2  # center horizontally

        try:
            pdf.image(img_path, x=x_pos, y=y_start, w=fit_w)
        except Exception as e:
            pdf.set_font(body_font, size=10)
            mcell(0, 6, f'[Graph render error: {e}]')

    # ── Section 3: Full Scripts (skip in preview mode) ──
    if not preview:
        pdf.add_page()
        pdf.set_font(body_font, size=18)
        cell_ln(0, 12, '◆ 第三部分: 完整剧本')
        pdf.ln(5)

        scripts = state.get('scripts', {})
        for ep in episodes:
            ep_id = ep.get('id', '')
            script_text = scripts.get(ep_id, '')
            if not script_text:
                continue

            pdf.add_page()
            pdf.set_font(body_font, size=16)
            cell_ln(0, 10, f"【{ep_id}】{ep.get('title', '')}")
            pdf.ln(3)

            # Replace glyphs missing from CJK font
            script_text = script_text.replace('\u26a0', '[!]')  # ⚠ → [!]
            script_text = script_text.replace('\u25b8', '\u25cf')  # ▸ → ●

            # Render script text with basic formatting
            for line in script_text.split('\n'):
                stripped = line.strip()
                if not stripped:
                    pdf.ln(3)
                    continue

                # Scene headers (new format: 场：NNN 景：xxx / 时：xxx 人：xxx)
                if re.match(r'^场：\d{3}', stripped) or stripped.startswith('时：'):
                    pdf.set_font(body_font, size=10)
                    mcell(0, 6, f"■ {stripped}")
                # Scene headers (legacy format)
                elif stripped.startswith('场景编号') or stripped.startswith('场景名称'):
                    pdf.set_font(body_font, size=10)
                    mcell(0, 6, f"■ {stripped}")
                # Stage directions
                elif stripped.startswith('▲') or stripped.startswith('△'):
                    pdf.set_font(body_font, size=9)
                    pdf.set_text_color(100, 100, 100)
                    mcell(0, 5, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Choice headers (new format: 选择 NNN)
                elif re.match(r'^选择\s+\d{3}', stripped):
                    pdf.ln(3)
                    pdf.set_font(body_font, size=11)
                    pdf.set_text_color(200, 80, 0)
                    mcell(0, 7, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Option lines (new format: NNNA：text)
                elif re.match(r'^\d{3}[A-D]：', stripped):
                    pdf.set_font(body_font, size=10)
                    pdf.set_text_color(100, 0, 150)
                    mcell(0, 6, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Branch markers (--- NNNA：xxx ---)
                elif re.match(r'^---\s+\d{3}[A-D]', stripped):
                    pdf.set_font(body_font, size=10)
                    pdf.set_text_color(100, 0, 150)
                    mcell(0, 6, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Choice blocks (legacy format)
                elif stripped.startswith('【互动节点'):
                    pdf.ln(3)
                    pdf.set_font(body_font, size=11)
                    pdf.set_text_color(200, 80, 0)
                    mcell(0, 7, stripped)
                    pdf.set_text_color(0, 0, 0)
                elif re.match(r'^[A-D]\.【', stripped):
                    pdf.set_font(body_font, size=10)
                    pdf.set_text_color(100, 0, 150)
                    mcell(0, 6, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Dead end (BE：)
                elif stripped.startswith('BE：') or stripped.startswith('BE:'):
                    pdf.set_font(body_font, size=10)
                    pdf.set_text_color(180, 0, 0)
                    mcell(0, 6, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Scene transitions & 判定 (green)
                elif re.match(r'^(\d{3}\s*场结束|选择\s+\d{3}[A-D]\s*结束)', stripped) or stripped.startswith('判定') or re.match(r'^若.+则接', stripped):
                    pdf.set_font(body_font, size=9)
                    pdf.set_text_color(0, 120, 0)
                    mcell(0, 5, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Chapter end markers
                elif re.match(r'^第.+[章集]\s*完', stripped):
                    pdf.ln(3)
                    pdf.set_font(body_font, size=11)
                    mcell(0, 7, stripped)
                # AI shot descriptions (cyan)
                elif stripped.startswith('AI ') or stripped.startswith('AI：') or stripped.startswith('AI:'):
                    pdf.set_font(body_font, size=9)
                    pdf.set_text_color(0, 140, 180)
                    mcell(0, 5, stripped)
                    pdf.set_text_color(0, 0, 0)
                # QTE markers
                elif stripped.startswith('【QTE'):
                    pdf.set_font(body_font, size=10)
                    pdf.set_text_color(200, 80, 0)
                    mcell(0, 6, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Character intro cards (amber)
                elif stripped.startswith('人名字幕条') or '字幕条）' in stripped:
                    pdf.set_font(body_font, size=10)
                    pdf.set_text_color(180, 120, 0)
                    mcell(0, 6, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Inline stat changes (好感度+1, 修为-5)
                elif re.match(r'^[^：:\n]+[\+\-]\d+\s*$', stripped) and not re.match(r'^[\u4e00-\u9fff]{1,6}[：:]', stripped):
                    pdf.set_font(body_font, size=9)
                    pdf.set_text_color(0, 100, 0)
                    mcell(0, 5, stripped)
                    pdf.set_text_color(0, 0, 0)
                # System popups (legacy)
                elif stripped.startswith('【系统'):
                    pdf.set_font(body_font, size=9)
                    pdf.set_text_color(0, 100, 0)
                    mcell(0, 5, stripped)
                    pdf.set_text_color(0, 0, 0)
                # Dialogue (character: line)
                elif re.match(r'^[\u4e00-\u9fff]{1,6}[：:]', stripped):
                    pdf.set_font(body_font, size=10)
                    mcell(0, 6, stripped)
                else:
                    pdf.set_font(body_font, size=10)
                    mcell(0, 6, stripped)

    # Save PDF
    safe_title = re.sub(r'[^\w\u4e00-\u9fff-]', '_', title)
    suffix = '_预览' if preview else ''
    pdf_path = os.path.join(output_dir, f'互动剧本_{safe_title}{suffix}.pdf')
    pdf.output(pdf_path)
    print(f'PDF: {pdf_path}')
    return pdf_path


def render_graphviz_graphs(state: dict, project_dir: str, preview: bool = False) -> list[tuple[str, str]]:
    """Render episode graphs + overview graph as PNG via Graphviz.

    Graph accuracy rules:
    - Sequential edges between consecutive main-line scenes (no dot suffix)
    - after_scene edges to choice diamonds
    - Choice edges directly to target scenes (no intermediate option nodes)
    - Convergence edges from branch scenes (dot suffix) to next main-line scene
    - Overview graph showing full multi-route story flow
    """
    import graphviz

    CJK_FONT = 'Heiti SC'
    structure = state.get('structure', {})
    episodes = structure.get('episodes', [])
    tmp_dir = os.path.join(project_dir, 'tmp_graphs')
    os.makedirs(tmp_dir, exist_ok=True)

    def is_branch(sid):
        """Scene ID has dot suffix (e.g., EP01-C4.A) = branch variant."""
        return '.' in sid.split('-')[-1]

    def scene_node_style(scene):
        """Return graphviz node attrs based on scene type."""
        stype = scene.get('type', '')
        if stype == 'dead_end':
            return dict(shape='box', fillcolor='#ffcdd2', color='#c62828',
                        fontcolor='#b71c1c', style='filled')
        if stype == 'ending':
            return dict(shape='doubleoctagon', fillcolor='#c8e6c9',
                        color='#2e7d32', fontcolor='#1b5e20', style='filled')
        if stype == 'converge':
            return dict(shape='box', fillcolor='#bbdefb', color='#1565c0',
                        style='filled,bold', penwidth='2')
        if stype == 'opening':
            return dict(shape='box', fillcolor='#e8eaf6', color='#283593',
                        style='filled,rounded')
        return dict(shape='box', fillcolor='#e3f2fd', color='#1565c0',
                    style='filled,rounded')

    results = []

    # ── Per-episode graphs (skip in preview mode) ─────────────────
    if not preview:
     for ep in episodes:
        ep_id = ep.get('id', 'EP00')
        scenes = ep.get('scenes', [])
        choices = ep.get('choices', [])

        g = graphviz.Digraph(ep_id, format='png', engine='dot')
        g.attr('graph', fontname=CJK_FONT, rankdir='TB', bgcolor='white',
               pad='0.4', nodesep='0.5', ranksep='0.7', dpi='150')
        g.attr('node', fontname=CJK_FONT, fontsize='10', style='filled')
        g.attr('edge', fontname=CJK_FONT, fontsize='8', color='#78909c')

        scene_map = {s['id']: s for s in scenes}
        scene_ids = [s['id'] for s in scenes]

        # Build lookups
        leads_to_targets = set()
        for ch in choices:
            for opt in ch.get('options', []):
                lt = opt.get('leads_to', '')
                if lt:
                    leads_to_targets.add(lt)

        after_scene_to_choice = {}
        for ch in choices:
            after = ch.get('after_scene', '')
            if after:
                after_scene_to_choice[after] = ch['id']

        # Add scene nodes (with truncated description)
        for scene in scenes:
            sid = scene['id']
            desc = scene.get('description', '')
            desc_short = desc[:35] + ('…' if len(desc) > 35 else '') if desc else ''
            label_parts = [f"{sid}: {scene.get('name', '')}"]
            if desc_short:
                label_parts.append(desc_short)
            label = '\n'.join(label_parts)
            g.node(sid, label, **scene_node_style(scene))

        # Add choice nodes (compact question label)
        for choice in choices:
            cid = choice['id']
            q = choice.get('question', '')
            if len(q) > 30:
                q = q[:28] + '…'
            g.node(cid, f"{cid}\n{q}", shape='diamond',
                   fillcolor='#fff3e0', color='#e65100', fontsize='9',
                   style='filled')

        # ── EDGES ──
        main_scenes = [sid for sid in scene_ids if not is_branch(sid)]

        # 1. after_scene → choice edges (ALL scenes, not just main-line)
        for sid in scene_ids:
            if sid in after_scene_to_choice:
                g.edge(sid, after_scene_to_choice[sid])

        # 2. Sequential main-line edges (skip after_scenes, already connected)
        for i, sid in enumerate(main_scenes):
            if sid not in after_scene_to_choice and i + 1 < len(main_scenes):
                g.edge(sid, main_scenes[i + 1])

        # 3. All-branch episodes (e.g., EP10 with parallel route endings)
        if not main_scenes:
            # Group by first letter after dot: C1.A→"A", C3.B2→"B"
            route_groups = {}
            for sid in scene_ids:
                parts = sid.split('-')[-1]
                if '.' in parts:
                    suffix_char = parts[parts.index('.') + 1]
                    route_groups.setdefault(suffix_char, []).append(sid)

            # Subgraph clusters for layout (side-by-side routes)
            for letter, sids in sorted(route_groups.items()):
                # Derive route label from first scene name
                first_name = scene_map.get(sids[0], {}).get('name', '')
                if '\u00b7' in first_name:  # · separator
                    rlabel = first_name.split('\u00b7')[0]
                else:
                    rlabel = f'Route {letter}'
                with g.subgraph(name=f'cluster_{letter}') as sub:
                    sub.attr(label=rlabel, style='dashed', color='#aaaaaa',
                             fontname=CJK_FONT, fontsize='11')
                    for sid in sids:
                        sub.node(sid)
                    # Place route-specific choice nodes in the cluster
                    for ch in choices:
                        ch_parts = ch['id'].split('-')[-1]
                        if '.' in ch_parts:
                            ch_letter = ch_parts[ch_parts.index('.') + 1]
                            if ch_letter == letter:
                                sub.node(ch['id'])

            # Sequential edges within each route group
            for letter, sids in route_groups.items():
                for i, sid in enumerate(sids):
                    if sid not in after_scene_to_choice:
                        if (i + 1 < len(sids)
                                and sids[i + 1] not in leads_to_targets):
                            g.edge(sid, sids[i + 1])

        # 4. Choice → target scene edges (direct, no option nodes)
        for choice in choices:
            cid = choice['id']
            for opt in choice.get('options', []):
                lt = opt.get('leads_to', '')
                label_text = f"{opt.get('label', '')}. {opt.get('text', '')}"
                if lt and lt in scene_map:
                    edge_kw = dict(label=label_text, fontsize='8')
                    if opt.get('dead_end'):
                        edge_kw.update(color='#c62828', fontcolor='#c62828',
                                       style='dashed')
                    else:
                        edge_kw.update(color='#6a1b9a', fontcolor='#6a1b9a')
                    g.edge(cid, lt, **edge_kw)
                elif lt:
                    # Cross-episode reference
                    exit_id = f"exit_{cid}_{opt.get('label', 'X')}"
                    g.node(exit_id, f"→ {lt}", shape='plaintext',
                           fontsize='9', fontcolor='#555555')
                    g.edge(cid, exit_id, label=label_text, fontsize='8',
                           color='#6a1b9a', style='dashed')

        # 5. Convergence: branch scenes → next main-line scene
        for i, sid in enumerate(scene_ids):
            s = scene_map[sid]
            if is_branch(sid) and s.get('type') not in ('dead_end', 'ending'):
                for j in range(i + 1, len(scene_ids)):
                    if not is_branch(scene_ids[j]):
                        g.edge(sid, scene_ids[j], style='dashed',
                               color='#90a4ae')
                        break

        # Render
        out_path = os.path.join(tmp_dir, ep_id)
        try:
            g.render(out_path, cleanup=True)
            results.append((ep_id, out_path + '.png'))
        except Exception as e:
            print(f'  Warning: Could not render graph for {ep_id}: {e}')

    # ── Overview graph ──────────────────────────────────────────────
    title = state.get('metadata', {}).get('title', '')
    overview = graphviz.Digraph('OVERVIEW', format='png', engine='dot')
    overview.attr('graph', fontname=CJK_FONT, rankdir='TB', bgcolor='white',
                  pad='0.5', nodesep='0.6', ranksep='0.8', dpi='200',
                  label=f'《{title}》全剧互动结构总览',
                  fontsize='20', labelloc='t')
    overview.attr('node', fontname=CJK_FONT, fontsize='11', style='filled')
    overview.attr('edge', fontname=CJK_FONT, fontsize='9')

    route_colors = {'A': '#2e7d32', 'B': '#1565c0', 'C': '#6a1b9a',
                    'D': '#c62828'}

    # Check for new episode_graph structure (stats-driven DAG)
    episode_graph = structure.get('episode_graph', {})
    endings = structure.get('endings', [])
    dead_ends = structure.get('dead_ends', [])

    if episode_graph and episode_graph.get('nodes'):
        # ── New: Stats-driven DAG overview ──
        nodes = episode_graph['nodes']
        node_map = {n['id']: n for n in nodes}

        # Build episode lookup for summaries
        ep_lookup = {ep['id']: ep for ep in episodes}

        # Color palette for stat-gated edges
        stat_colors = {}
        player_stats = structure.get('player_stats', [])
        palette = ['#2e7d32', '#1565c0', '#6a1b9a', '#c62828', '#e65100']
        for i, ps in enumerate(player_stats):
            stat_colors[ps['id']] = palette[i % len(palette)]

        # Add episode nodes
        for node in nodes:
            nid = node['id']
            ntitle = node.get('title', '')
            ntype = node.get('type', 'shared')
            # Build label: ID + short title only (no subtitle/summary to keep boxes compact)
            if ntitle and len(ntitle) > 12:
                ntitle = ntitle[:11] + '…'
            label = f"{nid}\n{ntitle}" if ntitle else nid

            if ntype == 'shared':
                kw = dict(fillcolor='#e8eaf6', color='#283593',
                          style='filled,rounded')
            elif ntype == 'gated':
                gate = node.get('gate') or {}
                stat_id = gate.get('stat', '')
                gc = stat_colors.get(stat_id, '#e65100')
                kw = dict(fillcolor='#fff3e0', color=gc,
                          style='filled,rounded,bold', penwidth='2')
                # Add gate condition to label
                op = gate.get('operator', '>')
                val = gate.get('value', '?')
                stat_name = stat_id
                for ps in player_stats:
                    if ps['id'] == stat_id:
                        stat_name = ps.get('name', stat_id)
                        break
                label += f"\n[{stat_name} {op} {val}]"
            elif ntype == 'branch':
                # Branch nodes — show with distinct color
                kw = dict(fillcolor='#fff3e0', color='#e65100',
                          style='filled,rounded')
            elif ntype == 'fallback':
                kw = dict(fillcolor='#f5f5f5', color='#9e9e9e',
                          style='filled,rounded,dashed')
            elif ntype == 'ending':
                kw = dict(fillcolor='#c8e6c9', color='#2e7d32',
                          style='filled,bold', penwidth='2')
            else:
                kw = dict(fillcolor='#e3f2fd', color='#1565c0',
                          style='filled,rounded')

            # First node gets special styling
            if nid == episode_graph.get('entry'):
                kw.update(fillcolor='#e8eaf6', color='#283593',
                          style='filled,rounded,bold', penwidth='2')

            overview.node(nid, label, shape='box', **kw)

        # Add edges from episode_graph.nodes[].next
        for node in nodes:
            nid = node['id']
            for nx in node.get('next', []):
                # Handle both string ("EP02") and object ({"target": "EP02", "condition": ...}) formats
                if isinstance(nx, str):
                    target = nx
                    cond = None
                else:
                    target = nx.get('target', '')
                    cond = nx.get('condition')
                if not target:
                    continue
                if cond:
                    stat_id = cond.get('stat', '')
                    op = cond.get('operator', '>')
                    val = cond.get('value', '?')
                    stat_name = stat_id
                    for ps in player_stats:
                        if ps['id'] == stat_id:
                            stat_name = ps.get('name', stat_id)
                            break
                    ec = stat_colors.get(stat_id, '#e65100')
                    overview.edge(nid, target,
                                  label=f"{stat_name} {op} {val}",
                                  color=ec, fontcolor=ec,
                                  penwidth='1.5', fontsize='9')
                else:
                    # Default/unconditional edge
                    overview.edge(nid, target,
                                  color='#333333', penwidth='1.5')

        # Build ending→episode map by scanning scenes with type 'ending'
        # and matching ending names to scene names
        ending_ep_map = {}  # ending_id → source episode_id
        for ep in episodes:
            for sc in ep.get('scenes', []):
                if sc.get('type') == 'ending':
                    sc_name = sc.get('name', '')
                    for ending in endings:
                        ename = ending.get('name', '')
                        if ename and (ename == sc_name or ename in sc_name or sc_name in ename):
                            ending_ep_map[ending['id']] = ep['id']

        # Also check: ending prerequisites mentioning an episode
        for ending in endings:
            eid = ending['id']
            if eid in ending_ep_map:
                continue
            prereq = ending.get('prerequisites', '')
            m_ep = re.search(r'(EP\d+[a-z]?)', prereq)
            if m_ep:
                ending_ep_map[eid] = m_ep.group(1)

        # Fallback: connect unmatched endings to terminal nodes (nodes with empty next)
        terminal_nodes = [n['id'] for n in nodes if not n.get('next')]
        unmatched_endings = [e for e in endings if e['id'] not in ending_ep_map]
        if terminal_nodes and unmatched_endings:
            # Distribute evenly among terminal nodes
            for i, ending in enumerate(unmatched_endings):
                ending_ep_map[ending['id']] = terminal_nodes[i % len(terminal_nodes)]

        # Add ending nodes and edges
        ending_colors = ['#2e7d32', '#1565c0', '#6a1b9a', '#c62828', '#e65100', '#00695c']
        for i, ending in enumerate(endings):
            eid = ending['id']
            if eid in node_map:
                continue
            overview.node(eid, ending.get('name', eid),
                          shape='doubleoctagon', fillcolor='#c8e6c9',
                          color='#2e7d32', fontsize='9', style='filled')
            source_ep = ending_ep_map.get(eid)
            if source_ep:
                ec = ending_colors[i % len(ending_colors)]
                overview.edge(source_ep, eid, color=ec,
                              style='dashed', penwidth='1.2')

        # Build dead_end→episode map by scanning scenes with type 'dead_end'
        de_ep_map = {}
        for ep in episodes:
            for sc in ep.get('scenes', []):
                if sc.get('type') == 'dead_end':
                    sc_name = sc.get('name', '')
                    for de in dead_ends:
                        dname = de.get('name', '')
                        if dname and (dname == sc_name or dname in sc_name or sc_name in dname):
                            de_ep_map[de['id']] = ep['id']

        # Also match by scene_id prefix
        for de in dead_ends:
            did = de['id']
            if did in de_ep_map:
                continue
            scene_id = de.get('scene_id', '')
            m_ep = re.match(r'(EP\d+[a-z]?)', scene_id)
            if m_ep:
                de_ep_map[did] = m_ep.group(1)

        # Add dead end nodes and edges
        for de in dead_ends:
            did = de['id']
            if did in node_map:
                continue
            overview.node(did, de.get('name', did),
                          shape='box', fillcolor='#ffcdd2', color='#c62828',
                          fontsize='9', style='filled')
            source_ep = de_ep_map.get(did)
            if source_ep:
                overview.edge(source_ep, did, color='#c62828',
                              style='dashed', fontsize='8',
                              label=de.get('name', '')[:10])

    else:
        # ── fork_points or legacy fork_structure overview ──
        fork_points = normalize_fork_points(structure)

        if fork_points:
            # Build overview from fork_points
            ep_lookup = {ep['id']: ep for ep in episodes}

            # Collect all thread episode IDs for coloring
            thread_ep_colors = {}  # ep_id → color
            thread_palette = ['#2e7d32', '#1565c0', '#6a1b9a', '#c62828', '#e65100',
                             '#00695c', '#4527a0', '#ad1457']
            color_idx = 0
            for fp in fork_points:
                for thread_name, thread_info in fp.get('threads', {}).items():
                    tc = thread_palette[color_idx % len(thread_palette)]
                    for eid in thread_info.get('episodes', []):
                        thread_ep_colors[eid] = tc
                    color_idx += 1

            convergence_eps = {fp['convergence_episode'] for fp in fork_points if fp.get('convergence_episode')}
            fork_choice_eps = set()
            for fp in fork_points:
                fc = fp.get('fork_choice', '')
                if fc:
                    # Find which episode contains this choice
                    for ep in episodes:
                        for ch in ep.get('choices', []):
                            if ch['id'] == fc:
                                fork_choice_eps.add(ep['id'])

            # Add episode nodes
            for ep in episodes:
                eid = ep['id']
                ep_title = ep.get('title', '')
                if ep_title and len(ep_title) > 12:
                    ep_title = ep_title[:11] + '…'
                label = f"{eid}\n{ep_title}" if ep_title else eid

                if eid == episodes[0]['id']:
                    kw = dict(fillcolor='#e8eaf6', color='#283593',
                              style='filled,rounded,bold', penwidth='2')
                elif eid in convergence_eps:
                    kw = dict(fillcolor='#bbdefb', color='#1565c0',
                              style='filled,bold', penwidth='2')
                elif eid in thread_ep_colors:
                    tc = thread_ep_colors[eid]
                    kw = dict(fillcolor='#fff3e0', color=tc,
                              style='filled,rounded,bold', penwidth='1.5')
                else:
                    kw = dict(fillcolor='#e3f2fd', color='#1565c0',
                              style='filled,rounded')
                overview.node(eid, label, shape='box', **kw)

            # Add ending/dead-end nodes
            for ending in endings:
                overview.node(ending['id'], ending.get('name', ending['id']),
                              shape='doubleoctagon', fillcolor='#c8e6c9',
                              color='#2e7d32', fontsize='9', style='filled')
            for de in dead_ends:
                overview.node(de['id'], de.get('name', de['id']),
                              shape='box', fillcolor='#ffcdd2', color='#c62828',
                              fontsize='9', style='filled')

            # Build edges from fork_points
            color_idx = 0
            for fp in fork_points:
                fork_choice = fp.get('fork_choice', '')
                convergence = fp.get('convergence_episode')

                # Find the episode containing the fork choice
                fork_ep = None
                if fork_choice:
                    for ep in episodes:
                        for ch in ep.get('choices', []):
                            if ch['id'] == fork_choice:
                                fork_ep = ep['id']
                                break
                        if fork_ep:
                            break

                for thread_name, thread_info in fp.get('threads', {}).items():
                    tc = thread_palette[color_idx % len(thread_palette)]
                    thread_eps = thread_info.get('episodes', [])

                    # Fork ep → first thread ep
                    if fork_ep and thread_eps:
                        overview.edge(fork_ep, thread_eps[0],
                                      label=thread_name, color=tc,
                                      fontcolor=tc, penwidth='1.5')

                    # Chain within thread
                    for i in range(len(thread_eps) - 1):
                        overview.edge(thread_eps[i], thread_eps[i + 1], color=tc)

                    # Last thread ep → convergence
                    if thread_eps and convergence:
                        overview.edge(thread_eps[-1], convergence,
                                      color=tc, style='dashed')

                    color_idx += 1

            # Add edges for non-fork sequential episodes
            all_thread_eps = set()
            for fp in fork_points:
                for tinfo in fp.get('threads', {}).values():
                    all_thread_eps.update(tinfo.get('episodes', []))

            shared_eps = [ep['id'] for ep in episodes if ep['id'] not in all_thread_eps]
            for i in range(len(shared_eps) - 1):
                # Only add edge if these are actually sequential (not separated by a fork)
                ep1_idx = next(j for j, ep in enumerate(episodes) if ep['id'] == shared_eps[i])
                ep2_idx = next(j for j, ep in enumerate(episodes) if ep['id'] == shared_eps[i + 1])

                # Check if there's a fork between them (thread eps in between)
                has_fork_between = False
                for j in range(ep1_idx + 1, ep2_idx):
                    if episodes[j]['id'] in all_thread_eps:
                        has_fork_between = True
                        break

                if not has_fork_between:
                    overview.edge(shared_eps[i], shared_eps[i + 1],
                                  color='#333333', penwidth='2')

            # Endings → find source episode via ending-type scenes
            ending_ep_map = {}
            for ep in episodes:
                for sc in ep.get('scenes', []):
                    if sc.get('type') == 'ending':
                        sc_name = sc.get('name', '')
                        for ending in endings:
                            ename = ending.get('name', '')
                            if ename and (ename == sc_name or ename in sc_name or sc_name in ename):
                                ending_ep_map[ending['id']] = ep['id']
            # Fallback: terminal episodes (not in any thread and at end, or last in thread with no convergence)
            terminal_eps = [ep['id'] for ep in episodes
                          if ep['id'] in all_thread_eps
                          and any(not fp.get('convergence_episode')
                                  for fp in fork_points
                                  for tinfo in fp.get('threads', {}).values()
                                  if ep['id'] in tinfo.get('episodes', []))]
            if not terminal_eps:
                terminal_eps = [episodes[-1]['id']]
            unmatched = [e for e in endings if e['id'] not in ending_ep_map]
            for i, ending in enumerate(unmatched):
                ending_ep_map[ending['id']] = terminal_eps[i % len(terminal_eps)]

            ending_colors = ['#2e7d32', '#1565c0', '#6a1b9a', '#c62828', '#e65100', '#00695c']
            for i, ending in enumerate(endings):
                eid = ending['id']
                source_ep = ending_ep_map.get(eid)
                if source_ep:
                    ec = ending_colors[i % len(ending_colors)]
                    overview.edge(source_ep, eid, color=ec, style='dashed',
                                  penwidth='1.2')

            # Dead ends → find source episode via dead_end-type scenes
            de_ep_map = {}
            for ep in episodes:
                for sc in ep.get('scenes', []):
                    if sc.get('type') == 'dead_end':
                        sc_name = sc.get('name', '')
                        for de in dead_ends:
                            dname = de.get('name', '')
                            if dname and (dname == sc_name or dname in sc_name or sc_name in dname):
                                de_ep_map[de['id']] = ep['id']
            for de in dead_ends:
                did = de['id']
                if did not in de_ep_map:
                    scene_id = de.get('scene_id', '')
                    m = re.match(r'(EP\d+[a-z]?)', scene_id)
                    if m:
                        de_ep_map[did] = m.group(1)
                source_ep = de_ep_map.get(did)
                if source_ep:
                    overview.edge(source_ep, did, color='#c62828',
                                  style='dashed', fontsize='8',
                                  label=de.get('name', '')[:10])

        else:
            # ── Lean format: build DAG from choice.options[].next ──
            ep_lookup = {ep['id']: ep for ep in episodes}

            # Detect thread colors from episode 'thread' field
            thread_colors = {}
            thread_palette = ['#333333', '#2e7d32', '#1565c0', '#6a1b9a',
                              '#c62828', '#e65100', '#00695c']
            seen_threads = []
            for ep in episodes:
                t = ep.get('thread', '')
                if t and t not in seen_threads:
                    seen_threads.append(t)
            for i, t in enumerate(seen_threads):
                thread_colors[t] = thread_palette[i % len(thread_palette)]

            # Add episode nodes
            for ep in episodes:
                eid = ep['id']
                ep_title = ep.get('title', '')
                if ep_title and len(ep_title) > 12:
                    ep_title = ep_title[:11] + '…'
                thread = ep.get('thread', '')
                thread_str = f"\n[{thread}]" if thread else ''
                label = f"{eid}\n{ep_title}{thread_str}" if ep_title else eid

                if eid == episodes[0]['id']:
                    kw = dict(fillcolor='#e8eaf6', color='#283593',
                              style='filled,rounded,bold', penwidth='2')
                elif ep.get('choice') is None:
                    # Terminal/ending episode
                    kw = dict(fillcolor='#c8e6c9', color='#2e7d32',
                              style='filled,bold', penwidth='2')
                else:
                    tc = thread_colors.get(thread, '#1565c0')
                    kw = dict(fillcolor='#e3f2fd', color=tc,
                              style='filled,rounded')

                overview.node(eid, label, shape='box', **kw)

            # Add edges from choice.options[].next
            for ep in episodes:
                eid = ep['id']
                ch = ep.get('choice')
                if not isinstance(ch, dict):
                    continue
                targets_seen = set()
                is_fork = ch.get('type') == 'fork'
                for opt in ch.get('options', []):
                    nxt = opt.get('next', '')
                    if not nxt or nxt in targets_seen:
                        continue
                    targets_seen.add(nxt)
                    edge_kw = dict(penwidth='1.5')
                    if is_fork:
                        # Color fork edges by target thread
                        target_ep = ep_lookup.get(nxt, {})
                        target_thread = target_ep.get('thread', '')
                        tc = thread_colors.get(target_thread, '#6a1b9a')
                        edge_kw.update(color=tc, fontcolor=tc,
                                       label=opt.get('text', ''))
                    else:
                        edge_kw.update(color='#333333')
                    overview.edge(eid, nxt, **edge_kw)

            # Add ending nodes and edges
            ending_colors = ['#2e7d32', '#1565c0', '#6a1b9a', '#c62828',
                             '#e65100', '#00695c']
            for i, ending in enumerate(endings):
                eid = ending['id']
                source_ep = ending.get('episode', '')
                if source_ep and source_ep in ep_lookup:
                    # Don't add duplicate node if ending ep is already a node
                    # Just add ending label node
                    overview.node(eid, ending.get('name', eid),
                                  shape='doubleoctagon', fillcolor='#c8e6c9',
                                  color='#2e7d32', fontsize='9', style='filled')
                    ec = ending_colors[i % len(ending_colors)]
                    overview.edge(source_ep, eid, color=ec, style='dashed',
                                  penwidth='1.2')

    out_path = os.path.join(tmp_dir, 'OVERVIEW')
    try:
        overview.render(out_path, cleanup=True)
        results.insert(0, ('OVERVIEW', out_path + '.png'))
    except Exception as e:
        print(f'  Warning: Could not render overview graph: {e}')

    return results


# ── DOCX rendering via python-docx ───────────────────────────────────

def render_docx(project_dir: str, state: dict):
    """Render individual DOCX files per episode."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    meta = state['metadata']
    title = meta.get('title', '互动影游')
    output_dir = os.path.join(project_dir, 'output')
    os.makedirs(output_dir, exist_ok=True)

    structure = state.get('structure', {})
    episodes = structure.get('episodes', [])
    scripts = state.get('scripts', {})
    safe_title = re.sub(r'[^\w\u4e00-\u9fff-]', '_', title)

    docx_paths = []
    for ep in episodes:
        ep_id = ep.get('id', '')
        script_text = scripts.get(ep_id, '')
        if not script_text:
            continue

        doc = Document()

        # Title
        p = doc.add_heading(f"{title} - {ep_id} {ep.get('title', '')}", level=0)
        if ep.get('subtitle'):
            doc.add_paragraph(ep['subtitle'])

        # Script content
        for line in script_text.split('\n'):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph('')
                continue

            if re.match(r'^场：\d{3}', stripped) or stripped.startswith('时：'):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
            elif stripped.startswith('场景编号') or stripped.startswith('场景名称'):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
            elif stripped.startswith('▲') or stripped.startswith('△'):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                run.font.size = Pt(9)
            elif re.match(r'^选择\s+\d{3}', stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(200, 80, 0)
            elif re.match(r'^\d{3}[A-D]：', stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.color.rgb = RGBColor(100, 0, 150)
            elif re.match(r'^---\s+\d{3}[A-D]', stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.color.rgb = RGBColor(100, 0, 150)
            elif stripped.startswith('【互动节点'):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(200, 80, 0)
            elif re.match(r'^[A-D]\.【', stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.color.rgb = RGBColor(100, 0, 150)
            elif stripped.startswith('BE：') or stripped.startswith('BE:'):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.color.rgb = RGBColor(180, 0, 0)
            elif re.match(r'^(\d{3}\s*场结束|选择\s+\d{3}[A-D]\s*结束)', stripped) or stripped.startswith('判定') or re.match(r'^若.+则接', stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.color.rgb = RGBColor(0, 120, 0)
                run.font.size = Pt(9)
            elif re.match(r'^第.+[章集]\s*完', stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(12)
            elif stripped.startswith('AI ') or stripped.startswith('AI：') or stripped.startswith('AI:'):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.italic = True
                run.font.color.rgb = RGBColor(0, 140, 180)
                run.font.size = Pt(9)
            elif stripped.startswith('【QTE'):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.color.rgb = RGBColor(200, 80, 0)
            elif stripped.startswith('人名字幕条') or '字幕条）' in stripped:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.color.rgb = RGBColor(180, 120, 0)
            elif re.match(r'^[^：:\n]+[\+\-]\d+\s*$', stripped) and not re.match(r'^[\u4e00-\u9fff]{1,6}[：:]', stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.color.rgb = RGBColor(0, 100, 0)
                run.font.size = Pt(9)
            elif stripped.startswith('【系统'):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.color.rgb = RGBColor(0, 100, 0)
                run.font.size = Pt(9)
            else:
                doc.add_paragraph(stripped)

        docx_path = os.path.join(output_dir, f'互动剧本_{safe_title}_{ep_id}.docx')
        doc.save(docx_path)
        docx_paths.append(docx_path)
        print(f'DOCX: {docx_path}')

    return docx_paths


# ── Main ─────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print('Usage: python renderer.py <project_dir> [--preview]')
        sys.exit(1)

    project_dir = sys.argv[1]
    preview = '--preview' in sys.argv
    state_path = os.path.join(project_dir, 'state.json')

    if not os.path.exists(state_path):
        print(f'Error: {state_path} not found')
        sys.exit(1)

    with open(state_path, 'r', encoding='utf-8') as f:
        state = json.load(f)

    print(f'Title: {state["metadata"].get("title", "?")}')
    print(f'Episodes: {len(state.get("structure", {}).get("episodes", []))}')
    print(f'Mode: {"preview" if preview else "full"}')
    print()

    pdf_path = render_pdf(project_dir, state, preview=preview)

    if not preview:
        docx_paths = render_docx(project_dir, state)
        print(f'\nDone. {1 + len(docx_paths)} files written.')
    else:
        print(f'\nDone. Preview PDF written.')


if __name__ == '__main__':
    main()
